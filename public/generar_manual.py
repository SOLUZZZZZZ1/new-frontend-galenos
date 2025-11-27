import os, re, html, datetime
from pathlib import Path

# --- Configuración ---
TITULO    = "Manual Profesional de Mediación y Métodos Alternativos de Resolución de Conflictos"
EDICION   = "Edición 2025"
AUTORES   = ["Mario Rondán Braida", "Sara Garrido García", "Joan Alís Gabernet", "Agustín Azparren Lucas"]
ENTIDAD   = "MEDIAZION – Centro Institucional de Mediación y Resolución de Conflictos"
LOGO      = "logo.png"    # asegúrate de tenerlo al lado del script
REEMPLAZO = ("IIMAT", "MEDIAZION")

# PRÓLOGO (con el primer párrafo de Mario)
PROLOGO = """La sociedad moderna, en sus ámbitos civiles y mercantiles, es una sociedad dinámica y, a la vez, conflictiva. Necesita herramientas aceptadas que le ayuden a que ese dinamismo no se atasque cuando un conflicto se produce. Los tribunales de justicia antaño ponían paz —o estaban para eso—, pero hoy resultan lentos y costosos para ciertas necesidades actuales. Por eso tienen cabida los Métodos Adecuados de Solución de Controversias (MASC), entre los que destaca la mediación. Es un método antiguo que, bien empleado, solucionará no pocos conflictos, descargando el sistema judicial y ofreciendo respuestas apropiadas a problemas que agilizarán las relaciones civiles y mercantiles. Por ello, es necesaria la formación de mediadores y la preparación de otros profesionales para que aprendan a emplear este método de resolución de controversias, pues la ley reconoce que un acta de mediación con avenencia tiene valor de sentencia y, en caso de incumplimiento, puede ser ejecutada judicialmente.

En MEDIAZION, creemos que la excelencia profesional se alcanza cuando el conocimiento técnico se combina con una ética firme y una mirada humana. Este Manual de Mediación y Métodos Alternativos de Resolución de Conflictos nace con la vocación de ofrecer una guía clara, rigurosa y práctica para quienes trabajan cada día por la paz social: profesionales del Derecho, mediadores, psicólogos, trabajadores sociales, funcionarios públicos, directivos y cualquier persona comprometida con la transformación positiva de los conflictos.

La obra que tiene en sus manos reúne fundamentos, metodologías y herramientas que dialogan con los estándares europeos y las mejores prácticas internacionales. Aborda la mediación como un proceso estructurado y flexible, analiza la conciliación y el arbitraje, ofrece modelos de actas y documentos esenciales, y explora los desafíos y oportunidades que plantean las tecnologías emergentes —incluida la inteligencia artificial— en el ecosistema de los MASC.

Que estas páginas sirvan como brújula para navegar los desafíos del presente y como impulso para continuar construyendo entornos más justos, colaborativos y humanos. Desde MEDIAZION —con la mirada puesta en el interés superior de las personas y el bien común— reafirmamos nuestro compromiso con una cultura del entendimiento que honre la dignidad, la diversidad y la paz.

Mario Rondán Braida
Director de MEDIAZION – Centro Institucional de Mediación y Resolución de Conflictos
"""

# Orden lógico de tus archivos (exactamente como me los pasaste)
ARCHIVOS = [
    "1.-DOCTRINA Y REGLAMENTOS DEL IIMAT copy copy.pdf",
    "2.-MASC ADR  definiciones y procedimientos ideas generales ok.doc",
    "3.-legislacion de mediacion BOE-A-2012-9112-consolidado (1) copy.pdf",
    "4.-CONCILIACIÓN IIMAT copy.pdf",
    "5.-MEDIACION REGLAMENTO.doc",
    "6.-mediacion electronica.doc",
    "7.-MODELOS DE ACTAS  DE MEDIACIÓN copy.pdf",
    "8-arbitraje.doc",
    "9.-TABLA DE ARANCELES ARBITRAJE Y MEDIACION.doc",
    "10.-La_inteligencia_artificial_y_los_MASC copy.pdf",
    "11.-El Derecho Colaborativo[1].doc",
    "12.-clausula IIMAT.doc",
    "13.-Corte del  IIMAT.doc",
]

# ----------------- utilidades -----------------
def limpiar(x:str)->str:
    if not x: return ""
    a = re.sub(r"\ufeff", "", x)
    a = a.replace("\r\n","\n").replace("\r","\n")
    a = re.sub(r"\n{3,}", "\n\n", a)
    # reemplazo IIMAT -> MEDIAZION
    a = a.replace(REEMPLAZO[0], REEMPLAZO[1])
    return a.strip()

def extraer_texto(ruta:Path)->str:
    suf = ruta.suffix.lower()
    try:
        if suf in [".docx"]:
            import docx
            d = docx.Document(str(ruta))
            return "\n".join(p.text for p in d.paragraphs)
        elif suf in [".doc"]:  # intentar vía docx si es docx renombrado; si no, aviso
            try:
                import docx
                d = docx.Document(str(ruta))
                return "\n".join(p.text for p in d.paragraphs)
            except Exception:
                return f"[ATENCIÓN] Abrir este .doc en Word y Guardar como .docx para extracción limpia.\n\nArchivo: {ruta.name}"
        elif suf == ".pdf":
            from pdfminer.high_level import extract_text
            return extract_text(str(ruta))
        else:
            return ruta.read_text(encoding="utf-8", errors="ignore")
    except Exception as e:
        return f"[ERROR] No se pudo extraer texto de {ruta.name}: {e}"

def h(s): return html.escape(s, quote=False)

# ----------------- construir HTML -----------------
def construir_html(secciones:list[str])->str:
    autores = " · ".join(AUTORES)
    hoy = datetime.date.today().year
    html_top = f"""<!doctype html>
<html lang="es">
<head>
<meta charset="utf-8"/>
<title>Manual-MEDIAZION-2025 (ES)</title>
<meta name="viewport" content="width=device-width, initial-scale=1"/>
<style>
  :root{{ --ink:#111; --muted:#444; --brand:#0f172a; }}
  body{{ font-family: Cambria, Georgia, "Times New Roman", serif; color:var(--ink); line-height:1.55; margin:0; }}
  .page{{ max-width: 900px; margin: 40px auto; padding: 0 28px; }}
  header.cover{{ text-align:center; margin-top:80px; }}
  .logo{{ width:200px; margin:0 auto 18px; display:block; }}
  h1{{ font-size:30px; margin:0 0 10px; }}
  h2{{ font-size:24px; margin:30px 0 8px; color:var(--brand); }}
  h3{{ font-size:19px; margin:18px 0 6px; color:#1f2937; }}
  .subtitle{{ font-style:italic; color:var(--muted); }}
  .authors, .credits{{ color:var(--muted); }}
  .authors{{ margin:18px 0 28px; }}
  .hr{{ height:1px; background:#e5e7eb; margin:28px 0; }}
  footer{{ color:#555; margin:32px 0 60px; }}
  a.anchor{{ display:block; position:relative; top:-90px; visibility:hidden; }}
  @media print {{ .page{{ margin:15mm auto; padding:0 }} }}
</style>
</head>
<body>
<div class="page">
  <header class="cover">
    <img class="logo" src="{h(LOGO)}" alt="MEDIAZION"/>
    <h1>{h(TITULO)}</h1>
    <div class="subtitle">{h(EDICION)}</div>
    <div class="authors"><strong>Autores:</strong><br/>{h(autores)}</div>
    <div class="credits"><em>{h(ENTIDAD)}</em></div>
  </header>
  <div class="hr"></div>
"""
    html_bottom = f"""
  <div class="hr"></div>
  <footer class="small">© {hoy} MEDIAZION · Centro Institucional de Mediación y Resolución de Conflictos</footer>
</div>
</body>
</html>"""
    cuerpo = "\n\n".join(secciones)
    return html_top + cuerpo + html_bottom

def ptag(texto):
    # envuelve en <p> preservando saltos
    texto = h(texto).replace("\n\n","</p><p>").replace("\n","<br/>")
    return f"<p>{texto}</p>"

def main():
    base = Path(".")
    secciones = []

    # Prólogo
    secciones.append('<a class="anchor" id="prologo"></a><section><h2>Prólogo</h2>' + ptag(PROLOGO) + "</section>")

    # Secciones y extracción
    mapping = [
      ("1. Fundamentos y doctrina", [
         ("1.1 Doctrina y reglamentos institucionales", ARCHIVOS[0]),
         ("1.2 MASC / ADR: definiciones y procedimientos", ARCHIVOS[1]),
      ]),
      ("2. Marco normativo", [
         ("2.1 Legislación española de mediación", ARCHIVOS[2]),
         ("2.2 Estándares y referencias europeas", None),
      ]),
      ("3. Procesos y técnicas", [
         ("3.1 Conciliación", ARCHIVOS[3]),
         ("3.2 Mediación – Reglamento operativo", ARCHIVOS[4]),
         ("3.3 Mediación electrónica", ARCHIVOS[5]),
         ("3.4 Arbitraje", ARCHIVOS[7]),
         ("3.5 Derecho colaborativo", ARCHIVOS[10]),
      ]),
      ("4. Modelos y herramientas", [
         ("4.1 Modelos de actas de mediación", ARCHIVOS[6]),
         ("4.2 Tablas de aranceles", ARCHIVOS[8]),
         ("4.3 Cláusula de sometimiento a MEDIAZION", ARCHIVOS[11]),
      ]),
      ("5. Innovación y tecnología", [
         ("5.1 Inteligencia artificial y MASC", ARCHIVOS[9]),
      ]),
      ("6. Organización institucional", [
         ("6.1 Corte de MEDIAZION", ARCHIVOS[12]),
      ]),
      ("Apéndice. Mediación y confianza cuántica", [
         ("", "APENDICE_AUTOMATICO")
      ])
    ]

    for titulo, sub in mapping:
        bloque = [f'<section><h2>{h(titulo)}</h2>']
        for subt, fname in sub:
            if subt: bloque.append(f'<h3>{h(subt)}</h3>')
            if fname == "APENDICE_AUTOMATICO":
                texto = """PQC (post-quantum cryptography), ODR, identidad digital europea (eIDAS 2) y EuroQCI marcarán el estándar de confianza en la resolución alternativa de disputas. MEDIAZION adopta una postura ética y técnica “post-quantum ready”, promoviendo seguridad, trazabilidad y transparencia en mediación electrónica."""
            elif fname:
                ruta = base / fname
                if not ruta.exists():
                    texto = f"[AVISO] No se encontró el archivo: {fname}"
                else:
                    bruto = extraer_texto(ruta)
                    texto = limpiar(bruto)
            else:
                texto = ""
            bloque.append(ptag(texto))
        bloque.append("</section>")
        secciones.append("\n".join(bloque))

    html_final = construir_html(secciones)
    Path("Manual-MEDIAZION-2025_ES.html").write_text(html_final, encoding="utf-8")

    # También un RTF simple (sin estilos complejos)
    rtf = "{\\rtf1\\ansi\\deff0\\fs22 " + \
          TITULO.replace("\\","\\\\") + " \\line " + \
          EDICION.replace("\\","\\\\") + " \\line " + \
          "Autores: " + " · ".join(AUTORES).replace("\\","\\\\") + " \\line \\line " + \
          ENTIDAD.replace("\\","\\\\") + " \\line \\line " + \
          PROLOGO.replace("\\","\\\\").replace("\n","\n\\line ") + " }"
    Path("Manual-MEDIAZION-2025_ES.rtf").write_text(rtf, encoding="utf-8")

    print("OK: Generados Manual-MEDIAZION-2025_ES.html y Manual-MEDIAZION-2025_ES.rtf")

if __name__ == "__main__":
    main()
